"""Readers for things that are not web pages: plain text, Markdown, PDF, DOCX.

Each one produces the same block structure as an HTML adapter, so everything
downstream — audio, timings, search, the reader — is unchanged.
"""

from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

from ..document import Article, Block, BlockKind, Section

#: A report nobody is skimming has fewer pages than this; a PDF beyond it is
#: extracted page by page, synchronously, inside the request.
MAX_PDF_PAGES = 2000
#: A .docx is a zip. Its own bytes can be small and still declare gigabytes
#: of XML inside it -- checked from the zip's own directory, before anything
#: is decompressed.
MAX_DOCX_UNCOMPRESSED = 200 * 1024 * 1024

#: A line that is really a heading, in Markdown or in plain prose. Matched
#: against one line at a time, so it needs no MULTILINE.
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
#: The same shape, anywhere in a document. Detection used _MD_HEADING itself,
#: which is anchored to the ends of the *string*, so a heading was only ever
#: found in text one line long — and every pasted Markdown document with no
#: list and no link was read as plain prose, hashes and all.
_MD_HEADING_ANY = re.compile(r"^#{1,6}\s+\S", re.M)
_MD_QUOTE = re.compile(r"^>\s?(.*)$")
_MD_LIST = re.compile(r"^\s*(?:[-*+]|\d+[.)])\s+(.*)$")
_MD_RULE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
_MD_FENCE = re.compile(r"^\s*(```|~~~)")

#: Inline Markdown that should not be spoken as punctuation.
_MD_INLINE = [
    (re.compile(r"!\[[^\]]*\]\([^)]*\)"), ""),          # images
    (re.compile(r"\[([^\]]+)\]\([^)]*\)"), r"\1"),      # links keep their text
    (re.compile(r"`([^`]+)`"), r"\1"),
    (re.compile(r"\*\*([^*]+)\*\*"), r"\1"),
    (re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)"), r"\1"),
    (re.compile(r"__([^_]+)__"), r"\1"),
    (re.compile(r"~~([^~]+)~~"), r"\1"),
]


class UnsupportedDocument(RuntimeError):
    pass


def strip_markdown(text: str) -> str:
    for pattern, replacement in _MD_INLINE:
        text = pattern.sub(replacement, text)
    return text.strip()


def looks_like_markdown(text: str) -> bool:
    head = text[:4000]
    return bool(
        _MD_HEADING_ANY.search(head) or re.search(r"^\s*[-*+]\s+", head, re.M) or "](" in head
    )


def article_from_text(
    text: str,
    title: str = "",
    markdown: bool | None = None,
    source: str = "Pasted text",
) -> Article:
    """Turn plain text or Markdown into an article.

    Blank lines separate paragraphs, which is how people actually write and
    paste. Markdown headings become sections; so does a short standalone line
    in plain text, which is how most people mark a section anyway.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if markdown is None:
        markdown = looks_like_markdown(text)

    lines = text.split("\n")

    def stands_alone(i: int) -> bool:
        """True when this line is its own block, blank above and below.

        Without this check the first line of any hard-wrapped paragraph looks
        like a heading: short, and with no terminal punctuation.
        """
        above = i == 0 or not lines[i - 1].strip()
        below = i + 1 >= len(lines) or not lines[i + 1].strip()
        return above and below

    sections: list[Section] = []
    current = Section(title="")
    paragraph: list[str] = []
    in_fence = False
    quoting = False

    def flush() -> None:
        nonlocal quoting
        joined = " ".join(paragraph).strip()
        kind = BlockKind.QUOTE if quoting else BlockKind.PARA
        paragraph.clear()
        quoting = False
        if joined:
            current.blocks.append(Block(kind=kind, text=strip_markdown(joined)))

    def new_section(heading: str) -> None:
        nonlocal current
        flush()
        if current.blocks:
            sections.append(current)
        current = Section(title=strip_markdown(heading))

    for index, raw in enumerate(lines):
        line = raw.rstrip()

        if markdown and _MD_FENCE.match(line):
            in_fence = not in_fence
            flush()
            continue
        if in_fence:
            continue

        if not line.strip():
            flush()
            continue
        if _MD_RULE.match(line):
            flush()
            continue

        heading = _MD_HEADING.match(line) if markdown else None
        if heading:
            new_section(heading.group(2))
            continue

        quote = _MD_QUOTE.match(line) if markdown else None
        if quote:
            # Run-on `>` lines are one quotation, the same way run-on plain
            # lines are one paragraph. Split, each got its own "Start quote…
            # End quote" spoken around it.
            body = strip_markdown(quote.group(1))
            if quoting and body:
                paragraph.append(body)
            elif body:
                flush()
                quoting = True
                paragraph.append(body)
            continue
        if quoting:
            flush()

        item = _MD_LIST.match(line)
        if item:
            flush()
            body = strip_markdown(item.group(1))
            if body:
                current.blocks.append(Block(kind=BlockKind.LIST_ITEM, text=body))
            continue

        # A short line standing alone between blank lines is a heading in
        # everything but name.
        if (
            not markdown
            and not paragraph
            and stands_alone(index)
            and len(line) < 70
            and not line.endswith((".", "?", "!", ",", ";", ":"))
        ):
            new_section(line)
            continue

        paragraph.append(line.strip())

    flush()
    if current.blocks:
        sections.append(current)

    if not title:
        title = _first_title(sections) or "Untitled note"

    article = Article(title=title, sections=sections, source=source, lang="en")
    article.sections = [s for s in article.sections if s.blocks]
    if article.sections and not article.sections[0].title:
        article.sections[0].title = article.title
    return article.renumber()


def _first_title(sections: list[Section]) -> str:
    for section in sections:
        if section.title:
            return section.title
    for section in sections:
        for block in section.blocks:
            words = block.text.split()
            return " ".join(words[:9]) + ("…" if len(words) > 9 else "")
    return ""


# --------------------------------------------------------------------------
# binary formats
# --------------------------------------------------------------------------


def text_from_pdf(data: bytes) -> str:
    """Extract text from a PDF, keeping paragraph breaks.

    pypdf returns hard-wrapped lines, so lines are rejoined into paragraphs:
    a break only counts where the previous line ends a sentence or is short.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise UnsupportedDocument(
            "PDF support needs the documents extra: uv sync --extra documents"
        ) from exc

    # A file named .pdf is not always a PDF, and a real one can be truncated.
    # pypdf raises its own errors for both, and they used to reach the browser
    # as a 500 rather than "that file could not be read".
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedDocument(f"that PDF could not be read: {exc}") from exc

    if len(reader.pages) > MAX_PDF_PAGES:
        raise UnsupportedDocument(
            f"that PDF has {len(reader.pages)} pages, over the {MAX_PDF_PAGES} limit"
        )

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue

    joined = "\n\n".join(pages)
    return _rewrap(joined)


def _rewrap(text: str) -> str:
    """Undo hard wrapping so paragraphs survive as paragraphs."""
    out: list[str] = []
    buffer: list[str] = []

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            if buffer:
                out.append(" ".join(buffer))
                buffer = []
            out.append("")
            continue

        # A hyphen at a line end is a split word, not punctuation.
        if buffer and buffer[-1].endswith("-"):
            buffer[-1] = buffer[-1][:-1] + line
            continue

        buffer.append(line)
        # A sentence end, or a short line, closes the paragraph.
        if line.endswith((".", "?", "!", '"', "”")) and len(line) < 65:
            out.append(" ".join(buffer))
            buffer = []

    if buffer:
        out.append(" ".join(buffer))

    collapsed = re.sub(r"\n{3,}", "\n\n", "\n".join(out))
    return collapsed.strip()


def text_from_docx(data: bytes) -> tuple[str, str]:
    """Extract text and the document title from a .docx.

    Word styles say which paragraphs are headings, so those are re-emitted as
    Markdown headings and the shared text reader handles the rest.
    """
    try:
        import docx
    except ImportError as exc:
        raise UnsupportedDocument(
            "DOCX support needs the documents extra: uv sync --extra documents"
        ) from exc

    # A .docx is a zip, and its central directory declares each entry's
    # uncompressed size without decompressing anything -- read that first, so
    # a small file that unpacks into gigabytes is refused before it does.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            declared = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise UnsupportedDocument(f"that Word file could not be read: {exc}") from exc
    if declared > MAX_DOCX_UNCOMPRESSED:
        raise UnsupportedDocument("that Word file is too large once decompressed")

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedDocument(f"that Word file could not be read: {exc}") from exc

    title = (document.core_properties.title or "").strip()

    lines: list[str] = []
    for paragraph in document.paragraphs:
        body = paragraph.text.strip()
        if not body:
            lines.append("")
            continue

        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "2"
            lines.append("")
            lines.append(f"{'#' * min(int(level), 6)} {body}")
            lines.append("")
        elif "quote" in style:
            lines.append(f"> {body}")
        elif "list" in style:
            lines.append(f"- {body}")
        else:
            lines.append(body)

    return "\n".join(lines), title


def article_from_file(data: bytes, filename: str, title: str = "") -> Article:
    """Dispatch on the extension. Raises for anything unhandled."""
    suffix = Path(filename).suffix.lower()
    stem = Path(filename).stem.replace("-", " ").replace("_", " ").strip()

    if suffix == ".pdf":
        article = article_from_text(text_from_pdf(data), title=title, markdown=False, source="PDF")
        if not title and article.title == "Untitled note":
            article.title = stem or article.title
        return article

    if suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise UnsupportedDocument(
                "Old .doc files are not supported. Save it as .docx and try again."
            )
        body, embedded = text_from_docx(data)
        article = article_from_text(body, title=title or embedded, markdown=True, source="Word document")
        # A heading in the document is a better title than the filename.
        if not (title or embedded) and article.title == "Untitled note":
            article.title = stem or article.title
        return article

    if suffix in (".md", ".markdown", ".mdown", ".txt", ".text", ""):
        body = data.decode("utf-8", errors="replace")
        markdown = suffix not in (".txt", ".text")
        article = article_from_text(
            body,
            title=title,
            markdown=markdown or None,
            source="Markdown" if markdown else "Text file",
        )
        # A heading inside the file beats the filename, as for PDF and DOCX.
        if not title and article.title == "Untitled note":
            article.title = stem or article.title
        return article

    raise UnsupportedDocument(f"textcast cannot read {suffix or 'that file type'} yet")
